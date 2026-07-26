import re
from copy import deepcopy
from io import BytesIO
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.forms import modelformset_factory
from django.http import HttpResponse
from django.shortcuts import redirect, render
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .forms import BenefitItemForm, FAQItemForm, FeatureItemForm, LandingPageContentForm, ProcessStepForm, TestimonialForm
from .models import BenefitItem, FAQItem, FeatureItem, LandingPageContent, ProcessStep, Testimonial
from .skill_map import DEFAULT_PROJECT_CATEGORIES, DEFAULT_SKILL_CATEGORIES, SKILL_KEYWORD_MAP


EMPTY_BUILDER_DATA = {
    'personal': {
        'full_name': '',
        'email': '',
        'phone': '',
        'linkedin': '',
        'github': '',
        'portfolio': '',
        'location': '',
    },
    'job_requirements': {
        'job_title': '',
        'role': '',
        'job_description': '',
        'keywords': '',
        'required_skills': '',
    },
    'skill_categories': [],
    'projects': [],
    'education': [],
    'experience': [],
    'certifications': [],
    'honors_awards': [],
}


def _seed_default_content():
    content = LandingPageContent.objects.first()
    if not content:
        content = LandingPageContent.objects.create()

    if not FeatureItem.objects.exists():
        FeatureItem.objects.bulk_create([
            FeatureItem(title='ATS optimized', description='Generate a clean CV that is easy for recruiting systems to read.', icon='⚡', order=1),
            FeatureItem(title='Job-targeted content', description='Shape your summary and skill list around the exact role you want.', icon='🎯', order=2),
            FeatureItem(title='Instant downloads', description='Export your final CV as PDF or DOCX in one click.', icon='📄', order=3),
        ])

    if not ProcessStep.objects.exists():
        ProcessStep.objects.bulk_create([
            ProcessStep(title='Share your details', description='Enter your profile, background, and target role.', order=1),
            ProcessStep(title='Match the role', description='Add job requirements and skills to align the content.', order=2),
            ProcessStep(title='Generate and download', description='Preview the ATS-ready result and export it instantly.', order=3),
        ])

    if not BenefitItem.objects.exists():
        BenefitItem.objects.bulk_create([
            BenefitItem(title='Keyword optimization', description='Use the job description to include relevant terms naturally.', icon='🔑', order=1),
            BenefitItem(title='Professional formatting', description='Create a structured CV with strong sections and easy reading.', icon='🧩', order=2),
            BenefitItem(title='No account required', description='No signup, no login, and no long-term data storage.', icon='🔒', order=3),
        ])

    if not Testimonial.objects.exists():
        Testimonial.objects.bulk_create([
            Testimonial(quote='The output looked polished and recruiter-friendly from the very first draft.', author='Aisha K.', role='Software Engineer', order=1),
            Testimonial(quote='I used it to tailor my CV for three roles in one afternoon.', author='David J.', role='Marketing Lead', order=2),
        ])

    if not FAQItem.objects.exists():
        FAQItem.objects.bulk_create([
            FAQItem(question='Is my data safe?', answer='Yes. We do not store your information beyond the current session and you can download your CV immediately.', order=1),
            FAQItem(question='Can I export to PDF or DOCX?', answer='Yes. The preview page includes both options for quick downloads.', order=2),
        ])

    return content


def _get_builder_data(request):
    data = request.session.get('builder_data')
    if not data:
        data = deepcopy(EMPTY_BUILDER_DATA)
        request.session['builder_data'] = data

    return data


def _extract_skill_categories(job_description='', required_skills='', extra_skills='', keywords=''):
    text = ' '.join([job_description, required_skills, extra_skills, keywords]).lower()

    # keep only alphanumeric, +, ., /, - and spaces
    text = re.sub(r'[^a-z0-9+\s./-]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    categories = []
    for category in DEFAULT_SKILL_CATEGORIES:
        matches = []
        seen = set()

        for raw_term, display_term in SKILL_KEYWORD_MAP.get(category, []):
            if not raw_term:
                continue

            # use word boundary regex for more accurate matching and escape special regex characters
            pattern = r'\b' + re.escape(raw_term) + r'\b'
            if re.search(pattern, text):
                if display_term not in seen:
                    seen.add(display_term)
                    matches.append(display_term)

        if matches:
            categories.append({'category': category, 'items': ', '.join(matches)})

    if not categories:
        categories = [
            {'category': 'Project Management & Soft Skills', 'items': 'Communication, Problem Solving'},
        ]

    return categories


def _generate_summary(data):
    title = data.get('job_requirements', {}).get('job_title') or data.get('job_requirements', {}).get('role') or ''
    job_description = data.get('job_requirements', {}).get('job_description', '')
    skills = []

    for category in data.get('skill_categories', []):
        items = [item.strip() for item in category.get('items', '').split(',') if item.strip()]
        skills.extend(items)

    top_skills = ', '.join(skills[:6]) if skills else ''
    
    # if user provided a custom summary in job requirements
    custom_summary = data.get('job_requirements', {}).get('custom_summary', '').strip()
    if custom_summary:
        return custom_summary
    
    # generate summary based on job description and skills
    if title and top_skills:
        return f"{title} professional with expertise in {top_skills}. Proven track record of delivering high-impact solutions through technical excellence and collaborative problem-solving. Passionate about leveraging technology to drive business value and innovation."

    elif title:
        return f"{title} professional with a strong foundation in relevant technologies and methodologies. Experienced in delivering quality results through analytical thinking and effective collaboration. Committed to continuous learning and professional growth."

    elif top_skills:
        return f"Skilled professional with expertise in {top_skills}. Demonstrated ability to solve complex problems and deliver results in fast-paced environments. Strong communicator and team player committed to excellence."

    return 'Motivated professional seeking to leverage technical skills and experience to contribute to organizational success. Adaptable learner with strong problem-solving abilities and a passion for continuous improvement.'


def _calculate_score(data):
    keywords = [item.strip().lower() for item in data['job_requirements'].get('keywords', '').split(',') if item.strip()]
    required_skills = [item.strip().lower() for item in data['job_requirements'].get('required_skills', '').split(',') if item.strip()]
    all_keywords = keywords + required_skills

    text = ' '.join([
        data['job_requirements'].get('job_title', ''),
        data['job_requirements'].get('job_description', ''),
        data['job_requirements'].get('role', ''),
        ' '.join([item.get('items', '') for item in data.get('skill_categories', [])]),
        ' '.join([project.get('title', '') + ' ' + project.get('description', '') for project in data.get('projects', [])]),
        ' '.join([edu.get('institution', '') + ' ' + edu.get('degree', '') + ' ' + edu.get('details', '') for edu in data.get('education', [])]),
        ' '.join([exp.get('company', '') + ' ' + exp.get('role', '') + ' ' + exp.get('responsibilities', '') for exp in data.get('experience', [])]),
    ]).lower()

    matched_keywords = []
    for keyword in all_keywords:
        if keyword and re.search(r'\b' + re.escape(keyword) + r'\b', text):
            matched_keywords.append(keyword)

    keyword_total = max(1, len(all_keywords))
    keyword_count = len(matched_keywords)
    score = 55 + round((keyword_count / keyword_total) * 30)

    if data['personal'].get('full_name') and data['personal'].get('email') and data['personal'].get('phone'):
        score += 8

    if data['job_requirements'].get('job_title') and data['job_requirements'].get('job_description'):
        score += 7

    if data.get('skill_categories'):
        score += 5

    if data.get('projects') or data.get('experience') or data.get('education'):
        score += 5

    score = min(score, 98)
    suggestions = [kw for kw in all_keywords if kw and not re.search(r'\b' + re.escape(kw) + r'\b', text)]

    return score, suggestions, keyword_count, keyword_total


def landing_page(request):
    content = _seed_default_content()
    context = {
        'content': content,
        'features': FeatureItem.objects.all(),
        'steps': ProcessStep.objects.all(),
        'benefits': BenefitItem.objects.all(),
        'testimonials': Testimonial.objects.all(),
        'faqs': FAQItem.objects.all(),
    }
    return render(request, 'cv_app/landing.html', context)


def builder(request):
    step = int(request.GET.get('step', 1))
    data = _get_builder_data(request)

    if request.method == 'POST':
        step = int(request.POST.get('step', step))

        if step == 1:
            data['personal'] = {
                'full_name': request.POST.get('full_name', '').strip(),
                'email': request.POST.get('email', '').strip(),
                'phone': request.POST.get('phone', '').strip(),
                'linkedin': request.POST.get('linkedin', '').strip(),
                'github': request.POST.get('github', '').strip(),
                'portfolio': request.POST.get('portfolio', '').strip(),
                'location': request.POST.get('location', '').strip(),
            }

        elif step == 2:
            data['job_requirements'] = {
                'job_title': request.POST.get('job_title', '').strip(),
                'role': request.POST.get('role', '').strip(),
                'job_description': request.POST.get('job_description', '').strip(),
                'keywords': request.POST.get('keywords', '').strip(),
                'required_skills': request.POST.get('required_skills', '').strip(),
                'custom_summary': request.POST.get('custom_summary', '').strip(),
            }
            data['skill_categories'] = _extract_skill_categories(
                data['job_requirements'].get('job_description', ''),
                data['job_requirements'].get('required_skills', ''),
                data['job_requirements'].get('required_skills', ''),
                data['job_requirements'].get('keywords', ''),
            )

        elif step == 3:
            skill_categories = []
            category_names = request.POST.getlist('skill_category[]')
            category_items = request.POST.getlist('skill_items[]')
            custom_categories = request.POST.getlist('skill_custom_category[]')

            for i, (cat, items) in enumerate(zip(category_names, category_items)):
                # if custom category is provided
                if i < len(custom_categories) and custom_categories[i].strip():
                    cat = custom_categories[i].strip()

                if cat.strip() or items.strip():
                    skill_categories.append({'category': cat.strip(), 'items': items.strip()})

            data['skill_categories'] = skill_categories if skill_categories else data.get('skill_categories', [])

        elif step == 4:
            project_titles = request.POST.getlist('project_title[]')
            project_categories = request.POST.getlist('project_category[]')
            project_descriptions = request.POST.getlist('project_description[]')
            source_code_links = request.POST.getlist('source_code_link[]')
            project_links = request.POST.getlist('project_link[]')

            data['projects'] = [
                {
                    'title': title.strip(),
                    'category': category.strip(),
                    'description': description.strip(),
                    'source_code_link': source_code.strip(),
                    'project_link': link.strip(),
                }

                for title, category, description, source_code, link in zip(
                    project_titles, project_categories, project_descriptions, source_code_links, project_links
                )
                if title.strip() or category.strip() or description.strip() or source_code.strip() or link.strip()
            ]

        elif step == 5:
            company_names = request.POST.getlist('experience_company[]')
            experience_roles = request.POST.getlist('experience_role[]')
            experience_periods = request.POST.getlist('experience_period[]')
            experience_responsibilities = request.POST.getlist('experience_responsibilities[]')

            data['experience'] = [
                {
                    'company': company.strip(),
                    'role': role.strip(),
                    'period': period.strip(),
                    'responsibilities': responsibilities.strip(),
                }

                for company, role, period, responsibilities in zip(
                    company_names, experience_roles, experience_periods, experience_responsibilities
                )
                if company.strip() or role.strip() or period.strip() or responsibilities.strip()
            ]

        elif step == 6:
            degrees = request.POST.getlist('education_degree[]')
            institutions = request.POST.getlist('education_institution[]')
            periods = request.POST.getlist('education_period[]')
            gpa_cgpas = request.POST.getlist('education_gpa_cgpa[]')
            gpa_maxes = request.POST.getlist('education_gpa_max[]')
            education_details = request.POST.getlist('education_details[]')

            data['education'] = [
                {
                    'degree': degree.strip(),
                    'institution': institution.strip(),
                    'period': period.strip(),
                    'gpa_cgpa': gpa_cgpa.strip(),
                    'gpa_max': gpa_max.strip(),
                    'details': details.strip(),
                }

                for degree, institution, period, gpa_cgpa, gpa_max, details in zip(
                    degrees, institutions, periods, gpa_cgpas, gpa_maxes, education_details
                )
                if degree.strip() or institution.strip() or period.strip() or gpa_cgpa.strip() or gpa_max.strip() or details.strip()
            ]

        elif step == 7:
            cert_names = request.POST.getlist('certification_name[]')
            cert_issuers = request.POST.getlist('certification_issuer[]')
            cert_dates = request.POST.getlist('certification_date[]')
            cert_expirations = request.POST.getlist('certification_expiration[]')
            cert_ids = request.POST.getlist('certification_credential_id[]')
            cert_urls = request.POST.getlist('certification_credential_url[]')

            data['certifications'] = [
                {
                    'name': name.strip(),
                    'issuer': issuer.strip(),
                    'date': date.strip(),
                    'expiration': expiration.strip(),
                    'credential_id': cred_id.strip(),
                    'credential_url': cred_url.strip(),
                }

                for name, issuer, date, expiration, cred_id, cred_url in zip(
                    cert_names, cert_issuers, cert_dates, cert_expirations, cert_ids, cert_urls
                )
                if name.strip() or issuer.strip() or date.strip()
            ]

        elif step == 8:
            honor_titles = request.POST.getlist('honor_title[]')
            honor_issuers = request.POST.getlist('honor_issuer[]')
            honor_dates = request.POST.getlist('honor_date[]')
            honor_descriptions = request.POST.getlist('honor_description[]')

            data['honors_awards'] = [
                {
                    'title': title.strip(),
                    'issuer': issuer.strip(),
                    'date': date.strip(),
                    'description': description.strip(),
                }

                for title, issuer, date, description in zip(
                    honor_titles, honor_issuers, honor_dates, honor_descriptions
                )
                if title.strip() or issuer.strip() or date.strip()
            ]

        request.session['builder_data'] = data

        if step < 9:
            return redirect(f'/create/?step={step + 1}')

        return redirect('/preview/')

    if not data.get('skill_categories') and data.get('job_requirements', {}).get('job_description'):
        data['skill_categories'] = _extract_skill_categories(
            data['job_requirements'].get('job_description', ''),
            data['job_requirements'].get('required_skills', ''),
            data['job_requirements'].get('required_skills', ''),
            data['job_requirements'].get('keywords', ''),
        )

        request.session['builder_data'] = data

    context = {
        'step': step,
        'data': data,
        'skill_categories': data.get('skill_categories', []),
        'default_skill_categories': DEFAULT_SKILL_CATEGORIES,
        'defaultProjectCategories': DEFAULT_PROJECT_CATEGORIES,
    }

    return render(request, 'cv_app/builder.html', context)


def preview(request):
    data = _get_builder_data(request)
    score, suggestions, keyword_count, keyword_total = _calculate_score(data)
    context = {
        'data': data,
        'score': score,
        'keyword_count': keyword_count,
        'keyword_total': keyword_total,
        'keyword_suggestions': suggestions[:6],
        'generated_summary': _generate_summary(data),
    }

    return render(request, 'cv_app/preview.html', context)


def download_pdf(request):
    data = _get_builder_data(request)
    generated_summary = _generate_summary(data)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()

    # Available width = 8.5 - 0.5 - 0.5 = 7.5 inches
    avail_width = 7.5 * inch

    center_style = ParagraphStyle(
        'Center', parent=styles['Normal'],
        fontName='Helvetica', fontSize=11, leading=14,
        textColor=colors.HexColor('#111827'), alignment=1
    )
    name_style = ParagraphStyle(
        'Name', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=18, leading=22,
        textColor=colors.HexColor('#1d4ed8'), alignment=1
    )
    heading_style = ParagraphStyle(
        'Heading', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=14, leading=17,
        textColor=colors.HexColor('#111827'), spaceAfter=6
    )
    body_style = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontName='Helvetica', fontSize=11, leading=14,
        textColor=colors.HexColor('#111827'), alignment=4
    )
    body_bold_style = ParagraphStyle(
        'BodyBold', parent=body_style,
        fontName='Helvetica-Bold'
    )
    small_style = ParagraphStyle(
        'Small', parent=styles['Normal'],
        fontName='Helvetica', fontSize=11, leading=14,
        textColor=colors.HexColor('#4b5563'), alignment=1
    )
    bullet_style = ParagraphStyle(
        'Bullet', parent=body_style,
        leftIndent=18, bulletIndent=6, spaceBefore=1, spaceAfter=1
    )
    project_link_style = ParagraphStyle(
        'ProjectLink', parent=styles['Normal'],
        fontName='Helvetica', fontSize=11, leading=14,
        textColor=colors.HexColor('#1d4ed8'), alignment=0
    )
    project_title_style = ParagraphStyle(
        'ProjectTitle', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=12, leading=15,
        textColor=colors.HexColor('#111827')
    )
    project_source_style = ParagraphStyle(
        'ProjectSource', parent=styles['Normal'],
        fontName='Helvetica', fontSize=12, leading=15,
        alignment=2, textColor=colors.HexColor('#1d4ed8')
    )

    story = []

    # ===== HEADER - All Center Aligned =====
    # Line 1: Name
    name = data['personal'].get('full_name') or 'Your Name'
    story.append(Paragraph(name, name_style))

    # Line 2: Phone | Email | Location | LinkedIn | GitHub | Portfolio (ALL on ONE line)
    contact_parts = []
    if data['personal'].get('phone'):
        contact_parts.append(data['personal'].get('phone'))
    if data['personal'].get('email'):
        contact_parts.append(data['personal'].get('email'))
    if data['personal'].get('location'):
        contact_parts.append(data['personal'].get('location'))

    # Add clickable profile links
    if data['personal'].get('linkedin'):
        url = data['personal'].get('linkedin')
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        contact_parts.append(f'<a href="{url}" color="#1d4ed8">LinkedIn</a>')
    if data['personal'].get('github'):
        url = data['personal'].get('github')
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        contact_parts.append(f'<a href="{url}" color="#1d4ed8">GitHub</a>')
    if data['personal'].get('portfolio'):
        url = data['personal'].get('portfolio')
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        contact_parts.append(f'<a href="{url}" color="#1d4ed8">Portfolio</a>')

    if contact_parts:
        story.append(Paragraph(' | '.join(contact_parts), center_style))

    story.append(Spacer(1, 6))

    # ===== PROFESSIONAL SUMMARY =====
    if generated_summary and generated_summary.strip():
        story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
        story.append(Paragraph(generated_summary, body_style))
        story.append(Spacer(1, 6))

    # ===== TECHNICAL SKILLS =====
    skill_cats = [c for c in data.get('skill_categories', []) if c.get('items', '').strip()]
    if skill_cats:
        story.append(Paragraph('TECHNICAL SKILLS', heading_style))
        for category in skill_cats:
            story.append(Paragraph(f"<b>{category.get('category')}</b>: {category.get('items')}", body_style))
        story.append(Spacer(1, 6))

    # ===== PROJECTS =====
    projects = [p for p in data.get('projects', []) if p.get('title', '').strip() or p.get('description', '').strip()]
    if projects:
        story.append(Paragraph('PROJECTS', heading_style))
        for project in projects:
            title = project.get('title', '').strip() or 'Project'
            source_code = project.get('source_code_link', '').strip()
            description = project.get('description', '').strip()
            project_link = project.get('project_link', '').strip()

            # Title line: Title LEFT, [Source Code] RIGHT (clickable) - use Table for perfect alignment
            if source_code:
                if not source_code.startswith(('http://', 'https://')):
                    source_code = 'https://' + source_code
                title_para = Paragraph(f'<b>{title}</b>', project_title_style)
                source_para = Paragraph(f'<a href="{source_code}" color="#1d4ed8">[Source Code]</a>', project_source_style)
                
                # Table with 2 columns: title (left), source code (right)
                table = Table([[title_para, source_para]], colWidths=[avail_width * 0.75, avail_width * 0.25])
                table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ]))
                story.append(table)
            else:
                story.append(Paragraph(f'<b>{title}</b>', project_title_style))

            # Description as bullet points (NO extra bullet before the list)
            if description:
                bullets = []
                for line in description.splitlines():
                    line = line.strip()
                    if line:
                        bullets.append(line)
                if not bullets:
                    bullets = [description]
                
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(b, bullet_style), bulletColor=colors.HexColor('#2563eb')) for b in bullets],
                        bulletType='bullet',
                    )
                )

            # Project link at bottom (left-aligned)
            if project_link:
                if not project_link.startswith(('http://', 'https://')):
                    project_link = 'https://' + project_link
                story.append(Paragraph(
                    f'Project Link: <a href="{project_link}" color="#1d4ed8">{project_link}</a>',
                    project_link_style
                ))

            story.append(Spacer(1, 4))
        story.append(Spacer(1, 6))

    # ===== EDUCATION =====
    education_list = [e for e in data.get('education', []) if e.get('degree', '').strip() or e.get('institution', '').strip()]
    if education_list:
        story.append(Paragraph('EDUCATION', heading_style))
        for edu in education_list:
            degree = edu.get('degree', '').strip()
            institution = edu.get('institution', '').strip()
            period = edu.get('period', '').strip()
            gpa = edu.get('gpa_cgpa', '').strip()
            gpa_max = edu.get('gpa_max', '').strip()
            details = edu.get('details', '').strip()

            if degree and institution:
                story.append(Paragraph(f'<b>{degree}</b> | {institution} | <b>Graduated</b>: {period}', body_style))
            elif degree:
                story.append(Paragraph(f'<b>{degree}</b> | <b>Graduated</b>: {period}', body_style))
            elif institution:
                story.append(Paragraph(f'{institution} | <b>Graduated</b>: {period}', body_style))

            if gpa or gpa_max:
                gpa_text = f"CGPA/GPA: {gpa}"
                if gpa_max:
                    gpa_text += f" / {gpa_max}"
                story.append(Paragraph(gpa_text, body_style))

            if details:
                story.append(Paragraph(details, body_style))

            story.append(Spacer(1, 4))
        story.append(Spacer(1, 6))

    # ===== EXPERIENCE =====
    experience_list = [e for e in data.get('experience', []) if e.get('company', '').strip() or e.get('role', '').strip()]
    if experience_list:
        story.append(Paragraph('EXPERIENCE', heading_style))
        for exp in experience_list:
            company = exp.get('company', '').strip()
            role = exp.get('role', '').strip()
            period = exp.get('period', '').strip()
            responsibilities = exp.get('responsibilities', '').strip()

            if company and role:
                story.append(Paragraph(f'<b>{role}</b> @ {company}', body_bold_style))
            elif role:
                story.append(Paragraph(f'<b>{role}</b>', body_bold_style))
            elif company:
                story.append(Paragraph(f'<b>{company}</b>', body_bold_style))

            if period:
                story.append(Paragraph(period, small_style))

            if responsibilities:
                resp_items = [item.strip() for item in responsibilities.split('\n') if item.strip()]
                if resp_items:
                    story.append(
                        ListFlowable(
                            [ListItem(Paragraph(item, bullet_style), bulletColor=colors.HexColor('#2563eb')) for item in resp_items],
                            bulletType='bullet',
                        )
                    )

            story.append(Spacer(1, 4))
        story.append(Spacer(1, 6))

    # ===== CERTIFICATIONS =====
    certs = [c for c in data.get('certifications', []) if c.get('name', '').strip() or c.get('issuer', '').strip()]
    if certs:
        story.append(Paragraph('CERTIFICATIONS', heading_style))
        for cert in certs:
            name = cert.get('name', '').strip()
            issuer = cert.get('issuer', '').strip()
            date = cert.get('date', '').strip()
            cred_id = cert.get('credential_id', '').strip()
            cred_url = cert.get('credential_url', '').strip()

            if name and issuer and date:
                story.append(Paragraph(f'<b>{name}</b> | {issuer} | <b>Issued</b>: {date}', body_style))
            elif name and issuer:
                story.append(Paragraph(f'<b>{name}</b> | {issuer}', body_style))
            elif name:
                story.append(Paragraph(f'<b>{name}</b>', body_style))

            # Credential ID and Verify on same line
            cred_parts = []
            if cred_id:
                cred_parts.append(f'<b>Credential ID</b>: {cred_id}')
            if cred_url:
                if not cred_url.startswith(('http://', 'https://')):
                    cred_url = 'https://' + cred_url
                cred_parts.append(f'<a href="{cred_url}" color="#1d4ed8">Verify</a>')
            if cred_parts:
                story.append(Paragraph(' | '.join(cred_parts), body_style))

            story.append(Spacer(1, 4))
        story.append(Spacer(1, 6))

    # ===== HONORS & AWARDS =====
    honors = [h for h in data.get('honors_awards', []) if h.get('title', '').strip() or h.get('issuer', '').strip()]
    if honors:
        story.append(Paragraph('HONORS & AWARDS', heading_style))
        for honor in honors:
            title = honor.get('title', '').strip()
            issuer = honor.get('issuer', '').strip()
            date = honor.get('date', '').strip()
            desc = honor.get('description', '').strip()

            if title and issuer and date:
                story.append(Paragraph(f'<b>{title}</b> — {issuer} | <b>Date</b>: {date}', body_style))
            elif title and issuer:
                story.append(Paragraph(f'<b>{title}</b> — {issuer}', body_style))
            elif title:
                story.append(Paragraph(f'<b>{title}</b>', body_style))

            if desc:
                story.append(Paragraph(desc, body_style))

            story.append(Spacer(1, 4))

    doc.build(story)
    pdf_data = buffer.getvalue()
    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="cv_guide_cv.pdf"'

    return response


def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph in python-docx."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}"><w:r><w:rPr><w:color w:val="1d4ed8"/><w:u w:val="single"/></w:rPr><w:t>{text}</w:t></w:r></w:hyperlink>')
    paragraph._p.append(hyperlink)


def download_docx(request):
    data = _get_builder_data(request)
    generated_summary = _generate_summary(data)

    document = Document()

    # Set 0.5 inch margins
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(14)

    # Helper to add a heading
    def add_section_heading(text):
        heading = document.add_heading(text, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            run.font.size = Pt(14)
            run.font.name = 'Calibri'
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.space_before = Pt(0)
        return heading

    def add_centered_paragraph(text, bold=False, size=Pt(11), color=RGBColor(0x14, 0x21, 0x3D)):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = size
        run.font.bold = bold
        run.font.color.rgb = color
        return p

    def add_justified_paragraph(text, bold=False, size=Pt(11), color=RGBColor(0x14, 0x21, 0x3D), space_after=Pt(0), space_before=Pt(0)):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = space_before
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = size
        run.font.bold = bold
        run.font.color.rgb = color
        return p

    def add_bullet_point(text):
        p = document.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.line_spacing = Pt(14)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
        return p

    def add_body_bold_then_normal(bold_text, normal_text, space_after=Pt(0)):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        run_b = p.add_run(bold_text)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(11)
        run_b.font.bold = True
        run_b.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
        run_n = p.add_run(normal_text)
        run_n.font.name = 'Calibri'
        run_n.font.size = Pt(11)
        run_n.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
        return p

    # ===== HEADER - All Center Aligned =====
    # Line 1: Name
    name = data['personal'].get('full_name') or 'Your Name'
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(name)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    # Line 2: Phone | Email | Location | LinkedIn | GitHub | Portfolio (ALL on ONE line)
    contact_parts = []
    if data['personal'].get('phone'):
        contact_parts.append(data['personal'].get('phone'))
    if data['personal'].get('email'):
        contact_parts.append(data['personal'].get('email'))
    if data['personal'].get('location'):
        contact_parts.append(data['personal'].get('location'))

    links = []
    if data['personal'].get('linkedin'):
        links.append(('LinkedIn', data['personal'].get('linkedin')))
    if data['personal'].get('github'):
        links.append(('GitHub', data['personal'].get('github')))
    if data['personal'].get('portfolio'):
        links.append(('Portfolio', data['personal'].get('portfolio')))

    # Add clickable profile links to contact parts
    for label, url in links:
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        contact_parts.append((label, url))

    if contact_parts:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        
        for i, item in enumerate(contact_parts):
            if i > 0:
                run = p.add_run(' | ')
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            
            if isinstance(item, tuple):
                # It's a link (label, url)
                label, url = item
                _add_hyperlink(p, label, url)
            else:
                # It's plain text (phone, email, location)
                run = p.add_run(item)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # ===== PROFESSIONAL SUMMARY =====
    if generated_summary and generated_summary.strip():
        add_section_heading('PROFESSIONAL SUMMARY')
        add_justified_paragraph(generated_summary, space_after=Pt(6))

    # ===== TECHNICAL SKILLS =====
    skill_cats = [c for c in data.get('skill_categories', []) if c.get('items', '').strip()]
    if skill_cats:
        add_section_heading('TECHNICAL SKILLS')
        for category in skill_cats:
            cat_name = category.get('category', '').strip()
            cat_items = category.get('items', '').strip()
            if cat_name and cat_items:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(14)
                run_b = p.add_run(cat_name + ': ')
                run_b.font.name = 'Calibri'
                run_b.font.size = Pt(12)
                run_b.font.bold = True
                run_b.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
                run_n = p.add_run(cat_items)
                run_n.font.name = 'Calibri'
                run_n.font.size = Pt(11)
                run_n.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
        add_justified_paragraph('', space_after=Pt(6))  # spacer

    # ===== PROJECTS =====
    projects = [p for p in data.get('projects', []) if p.get('title', '').strip() or p.get('description', '').strip()]
    if projects:
        add_section_heading('PROJECTS')
        for project in projects:
            title = project.get('title', '').strip() or 'Project'
            source_code = project.get('source_code_link', '').strip()
            description = project.get('description', '').strip()
            project_link = project.get('project_link', '').strip()

            # Title line: Title LEFT, [Source Code] RIGHT (clickable)
            if source_code:
                if not source_code.startswith(('http://', 'https://')):
                    source_code = 'https://' + source_code
                # Create a table with 2 columns for perfect alignment
                table = document.add_table(rows=1, cols=2)
                table.autofit = True
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Left cell - Title
                left_cell = table.cell(0, 0)
                left_para = left_cell.paragraphs[0]
                left_para.paragraph_format.space_after = Pt(0)
                left_para.paragraph_format.space_before = Pt(0)
                left_para.paragraph_format.line_spacing = Pt(15)
                run_title = left_para.add_run(title)
                run_title.font.name = 'Calibri'
                run_title.font.size = Pt(12)
                run_title.font.bold = True
                run_title.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
                
                # Right cell - Source Code
                right_cell = table.cell(0, 1)
                right_para = right_cell.paragraphs[0]
                right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                right_para.paragraph_format.space_after = Pt(0)
                right_para.paragraph_format.space_before = Pt(0)
                right_para.paragraph_format.line_spacing = Pt(15)
                _add_hyperlink(right_para, '[Source Code]', source_code)
                
                # Set column widths
                for row in table.rows:
                    row.cells[0].width = Inches(5.5)
                    row.cells[1].width = Inches(1.5)
            else:
                add_justified_paragraph(title, bold=True, size=Pt(12), space_after=Pt(0))

            # Description as bullet points (NO extra bullet before the list)
            if description:
                bullets = [line.strip() for line in description.splitlines() if line.strip()]
                if not bullets:
                    bullets = [description]
                for bullet in bullets:
                    add_bullet_point(bullet)

            # Project link at bottom (left-aligned)
            if project_link:
                if not project_link.startswith(('http://', 'https://')):
                    project_link = 'https://' + project_link
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = Pt(14)
                run = p.add_run('Project Link: ')
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
                _add_hyperlink(p, project_link, project_link)

            add_justified_paragraph('', space_after=Pt(4))
        add_justified_paragraph('', space_after=Pt(6))

    # ===== EDUCATION =====
    education_list = [e for e in data.get('education', []) if e.get('degree', '').strip() or e.get('institution', '').strip()]
    if education_list:
        add_section_heading('EDUCATION')
        for edu in education_list:
            degree = edu.get('degree', '').strip()
            institution = edu.get('institution', '').strip()
            period = edu.get('period', '').strip()
            gpa = edu.get('gpa_cgpa', '').strip()
            gpa_max = edu.get('gpa_max', '').strip()
            details = edu.get('details', '').strip()

            if degree and institution:
                add_body_bold_then_normal(f'{degree} | {institution} | Graduated: {period}', '', space_after=Pt(0))
            elif degree:
                add_body_bold_then_normal(f'{degree} | Graduated: {period}', '', space_after=Pt(0))
            elif institution:
                add_body_bold_then_normal(f'{institution} | Graduated: {period}', '', space_after=Pt(0))

            if gpa or gpa_max:
                gpa_text = f'CGPA/GPA: {gpa}'
                if gpa_max:
                    gpa_text += f' / {gpa_max}'
                add_justified_paragraph(gpa_text, space_after=Pt(0))

            if details:
                add_justified_paragraph(details, space_after=Pt(0))

            add_justified_paragraph('', space_after=Pt(4))
        add_justified_paragraph('', space_after=Pt(6))

    # ===== EXPERIENCE =====
    experience_list = [e for e in data.get('experience', []) if e.get('company', '').strip() or e.get('role', '').strip()]
    if experience_list:
        add_section_heading('EXPERIENCE')
        for exp in experience_list:
            company = exp.get('company', '').strip()
            role = exp.get('role', '').strip()
            period = exp.get('period', '').strip()
            responsibilities = exp.get('responsibilities', '').strip()

            if company and role:
                add_justified_paragraph(f'{role} @ {company}', bold=True, size=Pt(12), space_after=Pt(0))
            elif role:
                add_justified_paragraph(role, bold=True, size=Pt(12), space_after=Pt(0))
            elif company:
                add_justified_paragraph(company, bold=True, size=Pt(12), space_after=Pt(0))

            if period:
                add_justified_paragraph(period, size=Pt(11), color=RGBColor(0x4B, 0x55, 0x63), space_after=Pt(0))

            if responsibilities:
                resp_items = [item.strip() for item in responsibilities.split('\n') if item.strip()]
                for item in resp_items:
                    add_bullet_point(item)

            add_justified_paragraph('', space_after=Pt(4))
        add_justified_paragraph('', space_after=Pt(6))

    # ===== CERTIFICATIONS =====
    certs = [c for c in data.get('certifications', []) if c.get('name', '').strip() or c.get('issuer', '').strip()]
    if certs:
        add_section_heading('CERTIFICATIONS')
        for cert in certs:
            name = cert.get('name', '').strip()
            issuer = cert.get('issuer', '').strip()
            date = cert.get('date', '').strip()
            cred_id = cert.get('credential_id', '').strip()
            cred_url = cert.get('credential_url', '').strip()

            if name and issuer and date:
                add_justified_paragraph(f'{name} | {issuer} | Issued: {date}', bold=False, space_after=Pt(0))
            elif name and issuer:
                add_justified_paragraph(f'{name} | {issuer}', bold=False, space_after=Pt(0))
            elif name:
                add_justified_paragraph(name, bold=False, space_after=Pt(0))

            # Credential ID and Verify on same line
            cred_parts = []
            if cred_id:
                cred_parts.append(('Credential ID: ', cred_id, False))
            if cred_url:
                if not cred_url.startswith(('http://', 'https://')):
                    cred_url = 'https://' + cred_url
                cred_parts.append(('Verify', cred_url, True))

            if cred_parts:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = Pt(14)
                for i, (label, value, is_link) in enumerate(cred_parts):
                    if i > 0:
                        run = p.add_run(' | ')
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
                    if is_link:
                        _add_hyperlink(p, value, value)
                    else:
                        run_l = p.add_run(label)
                        run_l.font.name = 'Calibri'
                        run_l.font.size = Pt(11)
                        run_l.font.bold = True
                        run_l.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
                        run_v = p.add_run(value)
                        run_v.font.name = 'Calibri'
                        run_v.font.size = Pt(11)
                        run_v.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

            add_justified_paragraph('', space_after=Pt(4))
        add_justified_paragraph('', space_after=Pt(6))

    # ===== HONORS & AWARDS =====
    honors = [h for h in data.get('honors_awards', []) if h.get('title', '').strip() or h.get('issuer', '').strip()]
    if honors:
        add_section_heading('HONORS & AWARDS')
        for honor in honors:
            title = honor.get('title', '').strip()
            issuer = honor.get('issuer', '').strip()
            date = honor.get('date', '').strip()
            desc = honor.get('description', '').strip()

            if title and issuer and date:
                add_justified_paragraph(f'{title} — {issuer} | Date: {date}', space_after=Pt(0))
            elif title and issuer:
                add_justified_paragraph(f'{title} — {issuer}', space_after=Pt(0))
            elif title:
                add_justified_paragraph(title, space_after=Pt(0))

            if desc:
                add_justified_paragraph(desc, space_after=Pt(0))

            add_justified_paragraph('', space_after=Pt(4))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="cv_guide_cv.docx"'
    document.save(response)

    return response


@login_required(login_url='/cv-admin/login/')
def cv_admin_dashboard(request):
    content = LandingPageContent.objects.first() or LandingPageContent.objects.create()
    FeatureFormSet = modelformset_factory(FeatureItem, form=FeatureItemForm, extra=1)
    ProcessFormSet = modelformset_factory(ProcessStep, form=ProcessStepForm, extra=1)
    BenefitFormSet = modelformset_factory(BenefitItem, form=BenefitItemForm, extra=1)
    TestimonialFormSet = modelformset_factory(Testimonial, form=TestimonialForm, extra=1)
    FAQFormSet = modelformset_factory(FAQItem, form=FAQItemForm, extra=1)

    if request.method == 'POST':
        main_form = LandingPageContentForm(request.POST, instance=content)
        features_formset = FeatureFormSet(request.POST, queryset=FeatureItem.objects.all())
        steps_formset = ProcessFormSet(request.POST, queryset=ProcessStep.objects.all())
        benefits_formset = BenefitFormSet(request.POST, queryset=BenefitItem.objects.all())
        testimonials_formset = TestimonialFormSet(request.POST, queryset=Testimonial.objects.all())
        faqs_formset = FAQFormSet(request.POST, queryset=FAQItem.objects.all())

        if all([
            main_form.is_valid(),
            features_formset.is_valid(),
            steps_formset.is_valid(),
            benefits_formset.is_valid(),
            testimonials_formset.is_valid(),
            faqs_formset.is_valid(),
        ]):
            main_form.save()
            features_formset.save()
            steps_formset.save()
            benefits_formset.save()
            testimonials_formset.save()
            faqs_formset.save()
            messages.success(request, 'Landing page content updated.')

            return redirect('cv_admin_dashboard')

    else:
        main_form = LandingPageContentForm(instance=content)
        features_formset = FeatureFormSet(queryset=FeatureItem.objects.all())
        steps_formset = ProcessFormSet(queryset=ProcessStep.objects.all())
        benefits_formset = BenefitFormSet(queryset=BenefitItem.objects.all())
        testimonials_formset = TestimonialFormSet(queryset=Testimonial.objects.all())
        faqs_formset = FAQFormSet(queryset=FAQItem.objects.all())

    context = {
        'main_form': main_form,
        'features_formset': features_formset,
        'steps_formset': steps_formset,
        'benefits_formset': benefits_formset,
        'testimonials_formset': testimonials_formset,
        'faqs_formset': faqs_formset,
    }

    return render(request, 'cv_app/cv_admin.html', context)


def cv_admin_login(request):
    if request.user.is_authenticated:
        return redirect('cv_admin_dashboard')

    error = None
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        user = authenticate(request, username=username, password=password)

        if user is not None and user.is_staff:
            login(request, user)
            return redirect('cv_admin_dashboard')

        error = 'Please use a valid staff account.'

    return render(request, 'cv_app/admin_login.html', {'error': error})


def cv_admin_logout(request):
    logout(request)
    return redirect('cv_admin_login')